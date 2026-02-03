# -*- coding: utf-8 -*-
from docx import Document

def extract_all_content(docx_path, output_file):
    """Extract all text content from a Word document including tables."""
    doc = Document(docx_path)
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write("DOCUMENT CONTENT - PARAGRAPHS\n")
        f.write("=" * 80 + "\n")
        
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                f.write(f"[P{i}] {para.text}\n")
        
        f.write("\n" + "=" * 80 + "\n")
        f.write("DOCUMENT CONTENT - TABLES\n")
        f.write("=" * 80 + "\n")
        
        for table_idx, table in enumerate(doc.tables):
            f.write(f"\n--- TABLE {table_idx + 1} ---\n")
            for row_idx, row in enumerate(table.rows):
                cells_text = []
                for cell in row.cells:
                    text = cell.text.strip().replace('\n', ' | ')
                    if text:
                        cells_text.append(text)
                if cells_text:
                    f.write(f"Row {row_idx}: {' || '.join(cells_text)}\n")
    
    print(f"Content extracted to {output_file}")

if __name__ == "__main__":
    extract_all_content(r"MD - Quy Trinh IVFMD.docx", "extracted_content.txt")
